from __future__ import annotations

import json
from pathlib import Path
from typing import TYPE_CHECKING

from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from loguru import logger

from src.infra.config import load_search_config
from src.infra.paths import DATA_DIR
from src.llm.factory import resolve_api_key
from src.memory.embeddings import create_local_embeddings

if TYPE_CHECKING:
    from src.llm.providers import ProviderConfig

_VECTOR_DIR = DATA_DIR / "vectorstore"
_INDEX_META = _VECTOR_DIR / "index_meta.json"
_DOCS_DIR = DATA_DIR / "workspace" / "knowledge"
_current_provider: ProviderConfig | None = None


def set_rag_provider(provider: ProviderConfig | None) -> None:
    global _current_provider
    _current_provider = provider


def _rag_config() -> dict:
    return load_search_config().get("rag", {})


def _ensure_dirs() -> None:
    _VECTOR_DIR.mkdir(parents=True, exist_ok=True)
    _DOCS_DIR.mkdir(parents=True, exist_ok=True)


def _embedding_spec(cfg: dict | None = None) -> dict[str, str]:
    cfg = cfg or _rag_config()
    mode = cfg.get("embedding_mode", "local")
    if mode == "api":
        return {
            "backend": "api",
            "model": cfg.get("embedding_model", "text-embedding-3-small"),
        }
    return {
        "backend": "local",
        "model": cfg.get("local_embedding_model", "BAAI/bge-small-zh-v1.5"),
    }


def _create_api_embeddings(
    provider: ProviderConfig | None,
    cfg: dict,
) -> OpenAIEmbeddings:
    model = cfg.get("embedding_model", "text-embedding-3-small")
    api_key = "not-set"
    base_url = None
    if provider and provider.type == "openai_compatible":
        api_key = resolve_api_key(provider) or "not-set"
        base_url = provider.base_url
    kwargs: dict = {"model": model, "api_key": api_key}
    if base_url:
        kwargs["base_url"] = base_url
    return OpenAIEmbeddings(**kwargs)


def create_embeddings(provider: ProviderConfig | None = None) -> Embeddings:
    """创建 Embedding 实例。默认使用本地模型，不依赖 LLM Provider。"""
    cfg = _rag_config()
    mode = cfg.get("embedding_mode", "local")

    if mode == "local":
        spec = _embedding_spec(cfg)
        embeddings, _ = create_local_embeddings(spec["model"])
        return embeddings

    if mode == "api":
        return _create_api_embeddings(provider, cfg)

    # auto: 先尝试 API，失败则回退本地
    try:
        api_emb = _create_api_embeddings(provider, cfg)
        api_emb.embed_query("ping")
        logger.info("使用 API Embedding: {}", cfg.get("embedding_model"))
        return api_emb
    except Exception as exc:
        logger.warning("API Embedding 不可用，回退本地模型: {}", exc)
        spec = _embedding_spec({**cfg, "embedding_mode": "local"})
        embeddings, _ = create_local_embeddings(spec["model"])
        return embeddings


def _load_meta() -> dict:
    if not _INDEX_META.exists():
        return {"documents": [], "chunk_count": 0}
    with _INDEX_META.open(encoding="utf-8") as f:
        return json.load(f)


def _save_meta(meta: dict) -> None:
    _ensure_dirs()
    with _INDEX_META.open("w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def _embeddings_for_index(provider: ProviderConfig | None = None) -> Embeddings:
    """按索引元数据中的 backend 加载匹配的 Embedding。"""
    meta = _load_meta()
    stored = meta.get("embedding_backend")
    if stored == "api":
        return _create_api_embeddings(provider, _rag_config())
    if stored == "local":
        model = meta.get("embedding_model") or _rag_config().get(
            "local_embedding_model", "BAAI/bge-small-zh-v1.5"
        )
        embeddings, _ = create_local_embeddings(model)
        return embeddings
    return create_embeddings(provider)


def _load_vectorstore(provider: ProviderConfig | None = None):
    from langchain_community.vectorstores import FAISS

    index_path = _VECTOR_DIR / "faiss_index"
    if not index_path.exists():
        return None
    embeddings = _embeddings_for_index(provider)
    return FAISS.load_local(
        str(index_path),
        embeddings,
        allow_dangerous_deserialization=True,
    )


def _save_vectorstore(store) -> None:
    index_path = _VECTOR_DIR / "faiss_index"
    store.save_local(str(index_path))


def _read_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    raise ValueError(f"不支持的文件类型: {suffix}")


def ingest_files(
    paths: list[Path],
    provider: ProviderConfig | None = None,
) -> tuple[int, int]:
    """导入文档到向量库，返回 (文件数, 块数)。"""
    from langchain_community.vectorstores import FAISS

    _ensure_dirs()
    cfg = _rag_config()
    exts = set(cfg.get("supported_extensions", [".txt", ".md", ".pdf", ".docx"]))
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=int(cfg.get("chunk_size", 500)),
        chunk_overlap=int(cfg.get("chunk_overlap", 80)),
    )

    documents: list[Document] = []
    imported_files = 0
    meta = _load_meta()
    known = set(meta.get("documents", []))

    for path in paths:
        if path.is_dir():
            files = [p for p in path.rglob("*") if p.is_file() and p.suffix.lower() in exts]
        else:
            files = [path] if path.suffix.lower() in exts else []

        for file_path in files:
            key = str(file_path.resolve())
            try:
                text = _read_file(file_path)
                if not text.strip():
                    continue
                dest = _DOCS_DIR / file_path.name
                if not dest.exists() or dest.read_bytes() != file_path.read_bytes():
                    dest.write_bytes(file_path.read_bytes())
                chunks = splitter.split_text(text)
                for i, chunk in enumerate(chunks):
                    documents.append(
                        Document(
                            page_content=chunk,
                            metadata={"source": file_path.name, "chunk": i},
                        )
                    )
                known.add(key)
                imported_files += 1
            except Exception as exc:
                logger.warning("读取文件失败 {}: {}", file_path, exc)
                raise ValueError(f"无法读取文件 {file_path.name}: {exc}") from exc

    if not documents:
        return 0, 0

    spec = _embedding_spec(cfg)
    current_backend = meta.get("embedding_backend")
    if current_backend and current_backend != spec["backend"]:
        logger.warning(
            "Embedding 后端变更 ({} -> {})，重建索引",
            current_backend,
            spec["backend"],
        )
        meta["chunk_count"] = 0

    embeddings = create_embeddings(provider)
    existing = _load_vectorstore(provider) if meta.get("chunk_count", 0) > 0 else None
    if existing and meta.get("embedding_backend") == spec["backend"]:
        existing.add_documents(documents)
        store = existing
    else:
        store = FAISS.from_documents(documents, embeddings)

    _save_vectorstore(store)
    meta["documents"] = sorted(known)
    meta["chunk_count"] = meta.get("chunk_count", 0) + len(documents)
    meta["embedding_backend"] = spec["backend"]
    meta["embedding_model"] = spec["model"]
    _save_meta(meta)
    return imported_files, len(documents)


def search_knowledge_base(query: str, provider: ProviderConfig | None = None) -> str:
    provider = provider or _current_provider
    cfg = _rag_config()
    top_k = int(cfg.get("top_k", 4))
    store = _load_vectorstore(provider)
    if not store:
        return "知识库为空。请先在侧边栏「导入文档」上传 .txt / .md / .pdf / .docx 文件。"

    try:
        docs = store.similarity_search(query, k=top_k)
    except Exception as exc:
        logger.exception("知识库检索失败")
        return f"知识库检索失败: {exc}"

    if not docs:
        return "知识库中未找到与问题相关的内容。"

    lines = [f"检索到 {len(docs)} 条相关片段：", ""]
    for i, doc in enumerate(docs, 1):
        source = doc.metadata.get("source", "unknown")
        lines.append(f"--- 片段 {i}（来源: {source}）---")
        lines.append(doc.page_content.strip())
        lines.append("")
    return "\n".join(lines)


def get_knowledge_stats() -> dict:
    meta = _load_meta()
    spec = _embedding_spec()
    return {
        "document_count": len(meta.get("documents", [])),
        "chunk_count": meta.get("chunk_count", 0),
        "has_index": (_VECTOR_DIR / "faiss_index").exists(),
        "embedding_backend": meta.get("embedding_backend") or spec["backend"],
        "embedding_model": meta.get("embedding_model") or spec["model"],
    }
